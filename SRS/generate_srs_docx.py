"""
Generate Professional SRS Document for Monit Paper Agency
Outputs: .docx with charts, tables, and automation growth section
"""

import os
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
CHART_DIR = os.path.join(OUTPUT_DIR, "charts")
os.makedirs(CHART_DIR, exist_ok=True)

# ── Color Palette ──────────────────────────────────────────────
BRAND_DARK   = "#1B3A4B"
BRAND_MID    = "#2E86AB"
BRAND_LIGHT  = "#A3CEF1"
BRAND_ACCENT = "#E8630A"
BRAND_GREEN  = "#2D936C"
BRAND_RED    = "#C1292E"
GRAY_BG      = "#F0F4F8"
WHITE        = "#FFFFFF"

# ── Utility helpers ────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    color_hex = color_hex.lstrip("#")
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:space="0" '
            f'w:color="{val.get("color","000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def styled_table(doc, headers, rows, col_widths=None, header_color=BRAND_DARK):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = GRAY_BG if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = "Calibri"
            set_cell_shading(cell, bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(BRAND_DARK.lstrip("#"))
        run.font.name = "Calibri"
    return h

def add_para(doc, text, bold=False, size=10, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


# ── Chart Generation ───────────────────────────────────────────

def chart_style(ax, title, xlabel, ylabel):
    ax.set_title(title, fontsize=13, fontweight="bold", color=BRAND_DARK, pad=12)
    ax.set_xlabel(xlabel, fontsize=10, color="#333")
    ax.set_ylabel(ylabel, fontsize=10, color="#333")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(labelsize=9)
    ax.grid(axis="y", alpha=0.3)

def gen_chart_manpower(path):
    fig, ax = plt.subplots(figsize=(7, 4))
    departments = ["Sales\nFollow-up", "Procurement\nTracking", "Billing &\nInvoicing",
                    "Stock\nManagement", "Dispatch\nPlanning", "Reporting &\nMIS"]
    before = [4, 3, 3, 2, 2, 2]
    after  = [2, 1, 1, 1, 1, 0]
    x = np.arange(len(departments))
    w = 0.32
    bars1 = ax.bar(x - w/2, before, w, label="Before Automation", color=BRAND_RED, alpha=0.85, edgecolor="white")
    bars2 = ax.bar(x + w/2, after, w, label="After Automation", color=BRAND_GREEN, alpha=0.85, edgecolor="white")
    ax.bar_label(bars1, padding=2, fontsize=9, fontweight="bold")
    ax.bar_label(bars2, padding=2, fontsize=9, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(departments, fontsize=8.5)
    ax.set_ylim(0, 6)
    ax.yaxis.set_major_locator(mticker.MaxNLocator(integer=True))
    chart_style(ax, "Manpower Requirement — Before vs After Automation", "", "Number of People")
    ax.legend(fontsize=9, loc="upper right")
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)

def gen_chart_response_time(path):
    fig, ax = plt.subplots(figsize=(7, 4))
    channels = ["Phone\nInquiry", "WhatsApp\nInquiry", "Stock\nCheck", "Order\nConfirmation",
                "Dispatch\nStatus", "Payment\nStatus"]
    before = [30, 45, 20, 60, 120, 60]
    after  = [5, 2, 1, 5, 2, 1]
    x = np.arange(len(channels))
    w = 0.32
    bars1 = ax.bar(x - w/2, before, w, label="Before (Minutes)", color="#E07A5F", edgecolor="white")
    bars2 = ax.bar(x + w/2, after, w, label="After (Minutes)", color="#81B29A", edgecolor="white")
    ax.bar_label(bars1, padding=2, fontsize=8, fontweight="bold")
    ax.bar_label(bars2, padding=2, fontsize=8, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(channels, fontsize=8.5)
    chart_style(ax, "Response Time — Before vs After Automation", "", "Time (Minutes)")
    ax.legend(fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)

def gen_chart_error_reduction(path):
    fig, ax = plt.subplots(figsize=(7, 4))
    errors = ["Double Stock\nBooking", "Wrong Billing\n(Firm Mix-up)", "Missed\nFollow-ups",
              "Price\nErrors", "Dispatch\nDelays", "Credit Limit\nViolation"]
    before = [12, 8, 25, 10, 15, 7]
    after  = [0, 0, 2, 1, 3, 0]
    x = np.arange(len(errors))
    w = 0.32
    bars1 = ax.bar(x - w/2, before, w, label="Before (Incidents/Month)", color="#E63946", alpha=0.85, edgecolor="white")
    bars2 = ax.bar(x + w/2, after, w, label="After (Incidents/Month)", color="#2D936C", alpha=0.85, edgecolor="white")
    ax.bar_label(bars1, padding=2, fontsize=8, fontweight="bold")
    ax.bar_label(bars2, padding=2, fontsize=8, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(errors, fontsize=8.5)
    chart_style(ax, "Operational Error Reduction — Before vs After", "", "Incidents per Month")
    ax.legend(fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)

def gen_chart_revenue_growth(path):
    fig, ax = plt.subplots(figsize=(7, 4))
    quarters = ["Q1\n(Go-Live)", "Q2", "Q3", "Q4", "Q5\n(Year 2)", "Q6", "Q7", "Q8"]
    baseline = [100, 100, 100, 100, 100, 100, 100, 100]
    projected = [100, 108, 118, 128, 140, 152, 162, 175]
    ax.fill_between(range(len(quarters)), baseline, alpha=0.15, color=BRAND_RED)
    ax.fill_between(range(len(quarters)), projected, alpha=0.15, color=BRAND_GREEN)
    ax.plot(range(len(quarters)), baseline, "o--", color=BRAND_RED, linewidth=2, markersize=6, label="Without Automation (Baseline)")
    ax.plot(range(len(quarters)), projected, "o-", color=BRAND_GREEN, linewidth=2.5, markersize=7, label="With Automation (Projected)")
    for i, v in enumerate(projected):
        if v != 100:
            ax.annotate(f"+{v-100}%", (i, v), textcoords="offset points", xytext=(0, 10),
                       ha="center", fontsize=8, fontweight="bold", color=BRAND_GREEN)
    ax.set_xticks(range(len(quarters)))
    ax.set_xticklabels(quarters, fontsize=9)
    ax.set_ylim(80, 200)
    chart_style(ax, "Projected Revenue Growth Index (Base = 100)", "", "Revenue Index")
    ax.legend(fontsize=9, loc="upper left")
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)

def gen_chart_inquiry_conversion(path):
    fig, ax = plt.subplots(figsize=(7, 4))
    stages = ["Inquiries\nReceived", "Responded\n< 5 min", "Quote\nSent", "Order\nConfirmed", "Delivered\nSuccessfully"]
    before_vals = [100, 40, 30, 20, 18]
    after_vals = [100, 95, 80, 55, 52]
    x = np.arange(len(stages))
    w = 0.32
    bars1 = ax.bar(x - w/2, before_vals, w, label="Before Automation", color="#E07A5F", edgecolor="white")
    bars2 = ax.bar(x + w/2, after_vals, w, label="After Automation", color="#81B29A", edgecolor="white")
    ax.bar_label(bars1, fmt='%d%%', padding=2, fontsize=8, fontweight="bold")
    ax.bar_label(bars2, fmt='%d%%', padding=2, fontsize=8, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(stages, fontsize=8.5)
    ax.set_ylim(0, 120)
    chart_style(ax, "Inquiry-to-Delivery Conversion Funnel", "", "Percentage (%)")
    ax.legend(fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)

def gen_chart_daily_ops(path):
    fig, ax = plt.subplots(figsize=(7, 4))
    tasks = ["Stock\nUpdate", "Order\nEntry", "Invoice\nGeneration", "Follow-up\nCalls", "MIS\nReports", "Dispatch\nCoordination"]
    manual_hrs = [2.0, 3.0, 2.5, 3.0, 2.0, 1.5]
    auto_hrs   = [0.1, 0.5, 0.3, 1.0, 0.1, 0.5]
    x = np.arange(len(tasks))
    w = 0.32
    bars1 = ax.bar(x - w/2, manual_hrs, w, label="Manual (Hours/Day)", color="#BC4749", edgecolor="white")
    bars2 = ax.bar(x + w/2, auto_hrs, w, label="Automated (Hours/Day)", color="#386641", edgecolor="white")
    ax.bar_label(bars1, fmt='%.1f', padding=2, fontsize=8, fontweight="bold")
    ax.bar_label(bars2, fmt='%.1f', padding=2, fontsize=8, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(tasks, fontsize=8.5)
    chart_style(ax, "Daily Operational Hours — Manual vs Automated", "", "Hours per Day")
    ax.legend(fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


# ── Generate all charts ────────────────────────────────────────
CHART_MANPOWER       = os.path.join(CHART_DIR, "manpower.png")
CHART_RESPONSE       = os.path.join(CHART_DIR, "response_time.png")
CHART_ERRORS         = os.path.join(CHART_DIR, "error_reduction.png")
CHART_REVENUE        = os.path.join(CHART_DIR, "revenue_growth.png")
CHART_CONVERSION     = os.path.join(CHART_DIR, "inquiry_conversion.png")
CHART_DAILY_OPS      = os.path.join(CHART_DIR, "daily_ops.png")

print("Generating charts...")
gen_chart_manpower(CHART_MANPOWER)
gen_chart_response_time(CHART_RESPONSE)
gen_chart_error_reduction(CHART_ERRORS)
gen_chart_revenue_growth(CHART_REVENUE)
gen_chart_inquiry_conversion(CHART_CONVERSION)
gen_chart_daily_ops(CHART_DAILY_OPS)
print("Charts generated.")


# ══════════════════════════════════════════════════════════════
#  BUILD THE DOCUMENT
# ══════════════════════════════════════════════════════════════

doc = Document()

# ── Page setup ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# ── COVER PAGE ─────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SOFTWARE REQUIREMENTS SPECIFICATION")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor.from_string(BRAND_DARK.lstrip("#"))
run.bold = True
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Integrated Paper Trading Management System")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(BRAND_MID.lstrip("#"))
run.font.name = "Calibri"

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Monit Paper Sales Agency\nMonit Paper Associates")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor.from_string(BRAND_ACCENT.lstrip("#"))
run.bold = True
run.font.name = "Calibri"

doc.add_paragraph()
doc.add_paragraph()

cover_info = [
    ["Document Version", "1.0"],
    ["Date", "February 9, 2026"],
    ["Prepared For", "Monit Paper Sales Agency / Monit Paper Associates, Indore, MP"],
    ["Primary Contact", "Mr. Monit Jain — +91 81200 44474"],
    ["Status", "Draft — Pending Client Approval"],
]
t = doc.add_table(rows=len(cover_info), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(cover_info):
    c0 = t.rows[i].cells[0]
    c1 = t.rows[i].cells[1]
    c0.text = ""
    c1.text = ""
    r0 = c0.paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.name = "Calibri"
    r1 = c1.paragraphs[0].add_run(v)
    r1.font.size = Pt(10)
    r1.font.name = "Calibri"
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c0.width = Cm(5)
    c1.width = Cm(10)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor.from_string(BRAND_RED.lstrip("#"))
run.bold = True

doc.add_page_break()

# ── TABLE OF CONTENTS (placeholder) ───────────────────────────
add_heading_styled(doc, "Table of Contents", 1)
toc_items = [
    "1. Introduction",
    "2. Overall Description",
    "3. User Roles & Permissions",
    "4. System Architecture Overview",
    "5. Module 1 — Procurement & Purchase Management",
    "6. Module 2 — Inventory & Stock Management",
    "7. Module 3 — Sales & Order Management",
    "8. Module 4 — Customer-Facing App & AI Chatbot",
    "9. Module 5 — Planning Engine Integration",
    "10. Module 6 — Dispatch & Logistics Management",
    "11. Module 7 — Billing, Invoicing & Tally Integration",
    "12. Module 8 — Payment & Credit Management",
    "13. Module 9 — Quality, Claims & Returns Management",
    "14. Module 10 — Reporting, Analytics & Market Intelligence",
    "15. Module 11 — CRM & Customer Relationship Management",
    "16. Module 12 — Configuration & Master Data",
    "17. Non-Functional Requirements",
    "18. Business Rules Engine",
    "19. Integration Requirements",
    "20. Growth After Automation — Impact Analysis",
    "21. Phased Rollout Plan",
    "22. Appendices",
]
for item in toc_items:
    add_para(doc, item, size=10, space_after=3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "1. Introduction", 1)

add_heading_styled(doc, "1.1 Purpose", 2)
add_para(doc, (
    "This Software Requirements Specification (SRS) defines the functional and non-functional "
    "requirements for an Integrated Paper Trading Management System to be developed for "
    "Monit Paper Sales Agency and Monit Paper Associates. The system aims to digitize and automate "
    "the end-to-end paper trading operations — from procurement and inventory management to sales, "
    "dispatch, customer engagement, and reporting."
))

add_heading_styled(doc, "1.2 Scope", 2)
scope_items = [
    "Procurement from 10–12 authorized paper mills with TAT tracking and delivery management.",
    "Inventory management across multiple godowns and locations in Indore with real-time stock reservation.",
    "Sales order lifecycle — inquiry, quotation, order confirmation, fulfillment, and invoicing.",
    "Customer-facing AI chatbot/app for availability checks, carton planning, and order placement.",
    "Planning engine integration for carton dimension optimization and paper size recommendation.",
    "Dispatch and logistics including multi-point pickup, truck load planning, and bill interception.",
    "Tally ERP integration for accounting, billing, and payment synchronization.",
    "Reporting and analytics for business intelligence, follow-up tracking, and market insights.",
    "Credit and payment management with rule-based multi-level approval workflows.",
]
for item in scope_items:
    add_bullet(doc, item)

add_heading_styled(doc, "1.3 Business Context", 2)
add_para(doc, (
    "Monit Paper Agency is a 30+ year old paper and paperboard trading business based in Indore, "
    "Madhya Pradesh. The agency operates as an authorized distributor for major Indian paper mills "
    "(ITC Limited, NR Agarwal, Gayatri Shakti, Emami Paper Mills, Balkrishna Paper Mills, JK Paper, "
    "and others) and serves the packaging, digital printing, and food-grade packaging industries."
))
add_para(doc, "Two Firms:", bold=True)
add_bullet(doc, "Monit Paper Sales Agency — handles ITC Limited (PSPD) business.")
add_bullet(doc, "Monit Paper Associates — handles NR Agarwal and other mill business.")
add_para(doc, (
    "Both firms share the same customer base. The internal decision of which firm fulfills an order "
    "is made by the sales/management team based on the supplying mill."
))

add_para(doc, "Business Models:", bold=True)
add_bullet(doc, "Direct Indenting — customer orders, Monit places order with mill, material goes directly or via godown.")
add_bullet(doc, "Stock-and-Sell — Monit maintains stock; two sub-models:")
add_bullet(doc, "Bundles — pre-packed sheets from mills.", level=1)
add_bullet(doc, "Reel Cutting (Sheeter) — Monit has 2 sheeters to convert reels into custom sheet sizes.", level=1)

add_para(doc, "Key Problem Statement:", bold=True, color=BRAND_RED)
add_para(doc, (
    "Currently, procurement is managed entirely in Excel and Word documents. Stock tracking uses Google Sheets. "
    "Billing runs on Tally. There is no integrated system. The business owner spends significant time manually "
    "filtering Excel data for follow-ups and reporting. Multiple salesmen accessing the same stock pool leads to "
    "double-booking conflicts. The goal is to reduce people-dependency and increase system-dependency — "
    "\"the system should be the expert, not the person.\""
))

add_heading_styled(doc, "1.4 Definitions & Acronyms", 2)
acronyms = [
    ("GSM", "Grams per Square Meter — paper thickness/weight measurement"),
    ("FBB", "Folding Box Board"),
    ("SBS", "Solid Bleached Sulfate board"),
    ("Duplex Board", "Multi-layered paperboard (Grey Back / White Back variants)"),
    ("Kraft", "Strong brown paper used for packaging"),
    ("Deckle", "Width of the paper machine / reel width"),
    ("Reel", "Roll of paper as produced by the mill"),
    ("Sheeter", "Machine that cuts reels into flat sheets"),
    ("TAT", "Turn Around Time — time from order to material readiness"),
    ("PO / SO", "Purchase Order / Sales Order"),
    ("DN / CN", "Debit Note / Credit Note"),
    ("PDC", "Post-Dated Cheque"),
    ("Stock Lot", "Non-prime / excess material sold at discounted rates by mills"),
    ("Indent", "Direct order placed with mill on behalf of customer"),
    ("E-Way Bill", "Electronic waybill required under GST for goods movement"),
    ("PSPD", "Paperboards and Specialty Papers Division (of ITC)"),
    ("Mill Run", "A production batch/cycle at the paper mill"),
]
styled_table(doc, ["Term", "Definition"], acronyms, col_widths=[4, 13])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 2 — OVERALL DESCRIPTION
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "2. Overall Description", 1)

add_heading_styled(doc, "2.1 Product Perspective", 2)
add_para(doc, "The system is a new, standalone web and mobile application that integrates with:")
integrations = [
    "Tally ERP (existing) — for accounting, billing, and payment data.",
    "Planning Engine (existing Sheeter App) — for carton optimization and paper size recommendation.",
    "WhatsApp Business API — for customer notifications and chatbot interaction.",
    "Email systems — for mill communication and order confirmations.",
    "Mill portals (ITC, etc.) — for order status tracking (future scope).",
]
for item in integrations:
    add_bullet(doc, item)

add_heading_styled(doc, "2.2 User Classes", 2)
user_classes = [
    ("Owner / Director", "Monit, Papa", "Final approval authority, full dashboard, all reports", "Full Access"),
    ("Sales Team", "3+ salesmen", "Inquiry mgmt, quotations, customer interaction, orders", "Sales Module"),
    ("Purchase Team", "Procurement staff", "Mill orders, delivery tracking, procurement planning", "Procurement Module"),
    ("Accountant", "Accounting staff", "Tally entries, billing, stock entries, payments", "Billing & Accounts"),
    ("Godown Staff", "Warehouse workers", "Stock receiving, packing slip verification", "Inventory Module"),
    ("Sheeter Operator", "Machine operators", "Reel-to-sheet cutting jobs, trim waste", "Sheeter Module"),
    ("Customer", "External (via App)", "Availability check, inquiry, order, tracking", "Customer App"),
    ("AI Agent", "System automated", "Auto availability, rule checks, notifications", "Automated"),
]
styled_table(doc, ["User Class", "Who", "Description", "Access Level"], user_classes, col_widths=[3.5, 3, 6, 3.5])

add_heading_styled(doc, "2.3 Operating Environment", 2)
add_bullet(doc, "Web Application — Admin panel, dashboards, all internal operations (responsive web).")
add_bullet(doc, "Mobile App — Customer-facing with AI chatbot (Android priority, iOS later).")
add_bullet(doc, "WhatsApp Integration — Chatbot for customers who prefer WhatsApp.")
add_bullet(doc, "Backend — Cloud-hosted, API-driven architecture.")
add_bullet(doc, "Database — Relational database (PostgreSQL) with Redis cache for real-time operations.")

add_heading_styled(doc, "2.4 Key Constraints", 2)
constraints = [
    "Tally remains the primary accounting system; the new system must sync with it, not replace it.",
    "Customer-facing app must NEVER expose stock quantities — only availability status.",
    "Multiple salesmen accessing the same stock pool — real-time conflict resolution is critical.",
    "Mill portals have varying transparency levels (ITC is stage-wise; others may be minimal).",
    "System must handle both firms (Agency + Associates) with unified customer view but separate billing.",
]
for c in constraints:
    add_bullet(doc, c)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 3 — USER ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "3. User Roles & Permissions", 1)

add_heading_styled(doc, "3.1 Role-Permission Matrix", 2)
role_headers = ["Module", "Owner", "Sales", "Purchase", "Accountant", "Godown", "Customer"]
role_rows = [
    ("Dashboard (Full)", "RW", "R", "R", "R", "—", "—"),
    ("Customer Master", "RW", "R", "R", "R", "—", "—"),
    ("Mill/Supplier Master", "RW", "R", "RW", "R", "—", "—"),
    ("Inquiry Management", "RW", "RW", "R", "—", "—", "RW (own)"),
    ("Sales Orders", "RW", "RW", "R", "R", "—", "R (own)"),
    ("Order Approval", "RW", "—", "—", "—", "—", "—"),
    ("Purchase Orders", "RW", "R", "RW", "R", "—", "—"),
    ("Inventory", "RW", "R", "RW", "RW", "RW", "—"),
    ("Dispatch Planning", "RW", "R", "RW", "R", "RW", "—"),
    ("Billing & Invoicing", "RW", "R", "R", "RW", "—", "—"),
    ("Payment & Credit", "RW", "R", "—", "RW", "—", "—"),
    ("Reports", "RW", "R(own)", "R(own)", "R(own)", "—", "—"),
    ("Configuration", "RW", "—", "—", "—", "—", "—"),
]
styled_table(doc, role_headers, role_rows)
add_para(doc, "R = Read, RW = Read/Write, — = No Access", size=8, color="#666666")

add_heading_styled(doc, "3.2 Multi-Level Approval Configuration", 2)
add_bullet(doc, "Level 1 — System Auto-Approval: All rules pass (credit OK, stock available, customer qualified).")
add_bullet(doc, "Level 2 — Sales Person Review: For flagged cases (partial stock, minor rule violations).")
add_bullet(doc, "Level 3 — Owner/Director Approval: Cases exceeding limits, large orders, or escalated cases.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 4 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "4. System Architecture Overview", 1)

add_heading_styled(doc, "4.1 Architecture Layers", 2)
arch_rows = [
    ("Client Layer", "Web App (Admin), Mobile App (Customer), WhatsApp Chatbot, Email Notifications"),
    ("API Gateway", "Centralized request routing, authentication, rate limiting"),
    ("Application Layer", "Purchase, Inventory, Sales, Dispatch, Billing, Payment, CRM, AI/Chatbot, Rules Engine, Reporting, Notification services"),
    ("Integration Layer", "Tally Bridge (Sync API), Planning Engine API, WhatsApp Business API, E-Way Bill API, E-Invoice API"),
    ("Data Layer", "PostgreSQL (Primary DB), Redis (Cache/Real-time), S3/Blob (Documents & Images)"),
]
styled_table(doc, ["Layer", "Components"], arch_rows, col_widths=[4, 13])

add_heading_styled(doc, "4.2 Key Architectural Decisions", 2)
decisions = [
    "Real-time stock locking — When a salesperson initiates an order, stock is soft-locked to prevent conflicts.",
    "Event-driven notifications — All state changes trigger notifications to relevant stakeholders.",
    "Configurable rule engine — Business rules stored as data, not code. Owner adjusts without developers.",
    "Tally as accounting truth — New system handles workflow; Tally remains the ledger. Two-way sync.",
    "AI chatbot as controlled gateway — Never directly commits stock. Auto-processes within rules or escalates.",
]
for d in decisions:
    add_bullet(doc, d)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 5 — PROCUREMENT MODULE
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "5. Module 1 — Procurement & Purchase Management", 1)

add_heading_styled(doc, "5.1 Overview", 2)
add_para(doc, (
    "This module handles the complete lifecycle of purchasing paper/board from mills — from "
    "identifying the need (customer orders or stock replenishment) through order placement, "
    "delivery tracking, and material receipt with weight reconciliation."
))

add_heading_styled(doc, "5.2 Mill/Supplier Master", 2)
pr_master = [
    ("PR-001", "Maintain master list of all 10-12 supplier mills with contact, products, payment terms", "High"),
    ("PR-002", "Mill product catalog: paper types, GSM range, deckle widths, sizes, MOQs", "High"),
    ("PR-003", "Configurable TAT per mill-product combination", "High"),
    ("PR-004", "Multi-level TAT: mill default, category, specific paper, special-run", "High"),
    ("PR-005", "Mill-wise pricing: base rates, dealer margins, volume discounts, active schemes", "High"),
    ("PR-006", "Track mill incentive programs — quarterly targets, rebates, early payment discounts", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], pr_master, col_widths=[2, 12, 2])

add_heading_styled(doc, "5.3 Purchase Order Lifecycle", 2)
pr_po = [
    ("PR-010", "Purchase need from: (a) Customer SO requiring material, (b) Stock replenishment", "High"),
    ("PR-011", "PO form: mill, paper type, GSM, size, quantity, delivery date, customer ref, instructions", "High"),
    ("PR-012", "Auto-suggest optimal mill based on TAT, pricing, truck load optimization", "Medium"),
    ("PR-013", "Club loading suggestion for same-region mills (e.g., 4 NR units in Vapi)", "High"),
    ("PR-014", "Generate standardized PO PDF (replace Word → PDF manual process)", "High"),
    ("PR-015", "Email PO directly to mill from system", "High"),
    ("PR-016", "Pending Order List — all open POs not yet fully received", "High"),
    ("PR-017", "Track: order date, expected delivery, current status, follow-up history", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], pr_po, col_widths=[2, 12, 2])

add_heading_styled(doc, "5.4 Mill Order Tracking", 2)
add_para(doc, "Order status flow (configurable per mill):", bold=True)
add_para(doc, (
    "Order Placed → Order Confirmed by Mill → Included in Run → Roll/Reel Ready → "
    "Sheet Cut (if applicable) → Packed → Ready for Dispatch → Dispatched → In Transit → Received at Godown"
), size=9, color=BRAND_MID)
pr_track = [
    ("PR-020", "Support stage-wise status updates reflecting mill's own tracking (ITC model)", "High"),
    ("PR-021", "Manual status entry + future API/email parsing for automated updates", "High"),
    ("PR-022", "Automated follow-up reminders when delivery dates approaching or overdue", "High"),
    ("PR-023", "Delivery date change log — old/new dates with reason for every revision", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], pr_track, col_widths=[2, 12, 2])

add_heading_styled(doc, "5.5 Material Receipt & Reconciliation", 2)
pr_receipt = [
    ("PR-030", "Create 'Material in Transit' inventory entry on dispatch day (not arrival)", "High"),
    ("PR-031", "Record dispatch details: truck/lorry #, driver, bill copy ref, consignment weight", "High"),
    ("PR-032", "Godown receipt: actual weight, packing slip verification, visible damage/shortage", "High"),
    ("PR-033", "Weight reconciliation: mill weight vs actual. Flag if discrepancy > tolerance (0.5%)", "High"),
    ("PR-034", "Convert 'In Transit' to godown stock upon receipt, recording specific location", "High"),
    ("PR-035", "Full traceability: receipt linked to original PO", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], pr_receipt, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 6 — INVENTORY MODULE
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "6. Module 2 — Inventory & Stock Management", 1)

add_heading_styled(doc, "6.1 Five Stock Statuses", 2)
stock_statuses = [
    ("Available (AVL)", "Physical stock in godown, unallocated", "Yes — Immediate"),
    ("Reel-to-Sheet (RTS)", "Available as reel, needs sheeter cutting", "Yes — With cutting lead time"),
    ("Booked (BKD)", "Allocated to specific customer/job; can be freed if urgent", "Conditional — Needs approval"),
    ("Ordered (ORD)", "PO placed with mill, not yet dispatched", "Yes — With TAT lead time"),
    ("In Transit (TRN)", "Dispatched from mill, on the way", "Yes — With transit time"),
]
styled_table(doc, ["Status", "Description", "Available for Sale?"], stock_statuses, col_widths=[4, 8, 5])

add_heading_styled(doc, "6.2 Multi-Location Management", 2)
inv_loc = [
    ("IN-001", "Track stock across all locations: Sanwer Road, Rajwada, Kailash Park, future sites", "High"),
    ("IN-002", "Each entry: location, paper type, mill, GSM, size, quantity, status, lot/batch ref", "High"),
    ("IN-003", "Inter-location stock transfers with full audit trail", "Medium"),
    ("IN-004", "Consolidated view (all locations) and location-wise view", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], inv_loc, col_widths=[2, 12, 2])

add_heading_styled(doc, "6.3 Real-Time Stock Reservation (Salesman Conflict Prevention)", 2)
add_para(doc, (
    "CRITICAL: This solves the exact problem described in the MOM — where one salesman sold "
    "material already cut for another customer's order, leading to customer sentiment damage."
), bold=True, color=BRAND_RED, size=9)
inv_lock = [
    ("IN-010", "Soft-lock (reserve) stock immediately when salesperson initiates order", "Critical"),
    ("IN-011", "Configurable lock expiry (e.g., 2 hours) — auto-release if order not confirmed", "High"),
    ("IN-012", "Second salesperson sees 'Booked for [Customer]' with alternatives offered", "Critical"),
    ("IN-013", "Only Owner can override and reallocate booked stock", "High"),
    ("IN-014", "All booking/unbooking actions logged with timestamps for audit", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], inv_lock, col_widths=[2, 12, 2])

add_heading_styled(doc, "6.4 Reel & Sheeter Management", 2)
inv_reel = [
    ("IN-020", "Separate reel inventory with deckle width, diameter/weight, mill lot number", "High"),
    ("IN-021", "Sheeter job orders: input reel specs, target sheet size, quantity, expected output", "High"),
    ("IN-022", "On completion: deduct reel, add sheets, record trim waste", "High"),
    ("IN-023", "Track trim waste as sellable byproduct (type, quantity, value)", "Medium"),
    ("IN-024", "Optimization suggestions: when reel-cut can fulfill a sheet order with minimal waste", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], inv_reel, col_widths=[2, 12, 2])

add_heading_styled(doc, "6.5 Stock Visibility Rules", 2)
add_para(doc, "CRITICAL BUSINESS RULE:", bold=True, color=BRAND_RED)
inv_vis = [
    ("IN-040", "Customer-facing system NEVER displays stock quantities. Only: Available / Partially / X days / Not Available.", "Critical"),
    ("IN-041", "Even internal sharing excludes quantities — prevents turnover reverse-engineering", "Critical"),
    ("IN-042", "Owner can manually hide specific stock (strategic reservation)", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], inv_vis, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 7 — SALES MODULE
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "7. Module 3 — Sales & Order Management", 1)

add_heading_styled(doc, "7.1 Two Sales Channels", 2)
add_para(doc, "Channel 1 — Proactive (Mill-Run Driven):", bold=True)
add_bullet(doc, "Mill sends daily run report → Monit matches to customer database → shares on WhatsApp groups → inquiries come in.")
add_para(doc, "Channel 2 — Reactive (Customer-Initiated):", bold=True)
add_bullet(doc, "Customer sends inquiry via WhatsApp/phone/app/chatbot → sales team or AI checks availability → quotation → order.")

add_heading_styled(doc, "7.2 Inquiry Management", 2)
sl_inq = [
    ("SL-001", "Capture every inquiry: customer, paper type, GSM, size, quantity, urgency, source channel", "High"),
    ("SL-002", "Auto-assign to mapped salesperson (or round-robin if unmapped)", "High"),
    ("SL-003", "Auto-check stock availability and display result to salesperson", "High"),
    ("SL-004", "Track inquiry-to-order conversion rate and non-conversion reasons", "High"),
    ("SL-005", "SLA-based escalation: 5 min for app, 30 min for phone/WhatsApp", "High"),
    ("SL-006", "Mill-run broadcast responses auto-create linked inquiries", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], sl_inq, col_widths=[2, 12, 2])

add_heading_styled(doc, "7.3 Quotation & Pricing", 2)
sl_price = [
    ("SL-010", "Customer-wise rate cards: negotiated rates per paper type, GSM, mill brand", "High"),
    ("SL-011", "Auto-populate pricing for rate-locked customers", "High"),
    ("SL-012", "Manual rate with system showing: mill cost, standard margin, suggested price", "High"),
    ("SL-013", "Price approval if rate below configurable minimum margin", "Medium"),
    ("SL-014", "Mill price change alerts → flag affected customer rate cards for review", "Medium"),
    ("SL-015", "Separate stock lot pricing with margin tracking", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], sl_price, col_widths=[2, 12, 2])

add_heading_styled(doc, "7.4 Sales Order & Approval", 2)
sl_so = [
    ("SL-020", "Generate SO: customer, paper spec, quantity, rate, firm (Agency/Associates), delivery date", "High"),
    ("SL-021", "'Full and final' confirmation sent to customer with change window", "High"),
    ("SL-022", "Customer acknowledgment tracking (app/chatbot/manual)", "High"),
    ("SL-023", "Determine fulfillment source: godown stock / in-transit / reel-cut / new mill order", "High"),
    ("SL-024", "Auto-trigger purchase requisition if mill order needed", "High"),
    ("SL-025", "Auto-select billing firm based on supplying mill", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], sl_so, col_widths=[2, 12, 2])

add_heading_styled(doc, "7.5 Order Approval Workflow", 2)
add_para(doc, "The approval engine is the heart of the sales module:", bold=True)
sl_approve = [
    ("SL-030", "Every order passes approval check: stock, credit, overdue, qualification, custom rules", "Critical"),
    ("SL-031", "Auto-Approved: ALL rules pass → order proceeds to fulfillment automatically", "High"),
    ("SL-032", "Escalated: Rule fails → summary to approver: customer, stock, credit, failed rule(s)", "Critical"),
    ("SL-033", "Approver gets: Approve / Approve with Conditions / Hold / Reject (one-touch)", "High"),
    ("SL-034", "Rejection NEVER communicated as 'rejected' — diplomatic stock-related response only", "Critical"),
    ("SL-035", "Two-step approval for high-risk: credit utilization < 5% remaining", "Medium"),
    ("SL-036", "Situational override: 'Mill Pressure Mode' temporarily relaxes limits (time-bound, logged)", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], sl_approve, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 8 — CUSTOMER APP & AI CHATBOT
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "8. Module 4 — Customer-Facing App & AI Chatbot", 1)
add_para(doc, (
    "This is the \"game changer\" module — a customer-facing interface (mobile app + WhatsApp chatbot) "
    "that acts as an AI-powered salesperson. It handles availability checks, suggests alternatives, "
    "collects inquiries, and routes orders through the approval system. Crucially, it must behave "
    "diplomatically and never expose internal business decisions to the customer."
), bold=True, size=9, color=BRAND_MID)

add_heading_styled(doc, "8.1 Availability Search", 2)
ca_search = [
    ("CA-001", "Search by paper type, GSM, sheet size, mill brand", "High"),
    ("CA-002", "Results show ONLY: Available (green), Partial (yellow), X days (blue), Not Available (red)", "Critical"),
    ("CA-003", "Available in X days = mill TAT + transit time", "High"),
    ("CA-004", "Alternative suggestions: nearby GSM, sizes, different mills, reel-cut options", "High"),
    ("CA-005", "Log every search for market intelligence (what, when, who, result, follow-up action)", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], ca_search, col_widths=[2, 12, 2])

add_heading_styled(doc, "8.2 Carton Dimension Optimizer", 2)
ca_plan = [
    ("CA-010", "Input: box L×W×H + box type (reverse tuck-in, lock bottom, etc. — 20-40 types)", "High"),
    ("CA-011", "With machine size on profile → Planning Engine calculates machine-specific optimal size", "High"),
    ("CA-012", "Without machine size → calculate based on standard stock sizes", "High"),
    ("CA-013", "Output: sheet size, ups, wastage %, stock availability for that size", "High"),
    ("CA-014", "Multiple plans ranked: least wastage, cheapest, fastest availability", "High"),
    ("CA-015", "Die match alert: 'Compatible die exists — saves ₹X on die cost'", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], ca_plan, col_widths=[2, 12, 2])

add_heading_styled(doc, "8.3 AI Chatbot — Conversational Sales", 2)
add_para(doc, "The chatbot must behave like a knowledgeable salesperson, not just a search tool:", bold=True)
ca_chat = [
    ("CA-020", "Conversational: understand context, suggest alternatives proactively", "High"),
    ("CA-021", "Example: '23x36 not in stock, but 23.5x36 available. NR has exact. ITC fresh run in 4 days.'", "High"),
    ("CA-022", "Machine-aware: 'Your machine takes 25x36. But 18x23 gives 12 ups/2% waste vs 6 ups/8%'", "High"),
    ("CA-023", "Reel-cut suggestion: 'Not in sheets but can cut from reel. Ready by [time].'", "Medium"),
    ("CA-024", "Stock lot proactive: 'Also have 210 GSM stock lot — can bill as 200 GSM at ₹5/kg discount'", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], ca_chat, col_widths=[2, 12, 2])

add_heading_styled(doc, "8.4 Hidden Rule Engine — Diplomatic Decision Flow", 2)
add_para(doc, (
    "The chatbot transparently serves the customer while invisibly enforcing business rules. "
    "The customer never sees the internal decision-making."
), size=9)

# Diplomatic response table
diplo_headers = ["Internal Situation", "Customer-Facing Response"]
diplo_rows = [
    ("Stock available, all rules pass", "\"Available! Would you like to order?\""),
    ("Stock available, credit issue (soft)", "\"Let me check if this is already booked… one moment.\" → Internal alert to owner"),
    ("Stock available, credit issue (hard)", "\"Let me connect you with our sales team for this one.\""),
    ("Stock not available", "\"Not available currently. Here are alternatives: [list]\""),
    ("Partially available", "\"Partially available now. Rest in [X] days.\""),
    ("Customer blocked", "\"I'll have the team reach out to you shortly.\""),
    ("Owner approves (from escalation)", "\"Great news — it's available! Shall I proceed?\""),
    ("Owner rejects (from escalation)", "\"This batch is committed. Sales team will contact you with options.\""),
]
styled_table(doc, diplo_headers, diplo_rows, col_widths=[7, 10], header_color=BRAND_ACCENT)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 9 — PLANNING ENGINE
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "9. Module 5 — Planning Engine Integration", 1)
add_para(doc, (
    "An existing Planning Engine (Sheeter App) is already in production with clients like Vijay Shri Printers. "
    "It calculates optimal sheet sizes from carton dimensions, considering machine sizes, die inventory, "
    "and wastage optimization. Integration is via REST API — planning only, not costing."
))
pe_reqs = [
    ("PE-001", "Integrate via API — planning only (not costing, which is printer-specific)", "High"),
    ("PE-002", "Customer dimension input → API call → display optimized sheet recommendations", "High"),
    ("PE-003", "Cross-reference results with current stock availability", "High"),
    ("PE-004", "Factor in customer's machine size if on profile", "High"),
    ("PE-005", "Store planning history for recommendation improvement", "Medium"),
    ("PE-006", "'What if' mode — modify parameters and see instant plan changes", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], pe_reqs, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 10 — DISPATCH & LOGISTICS
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "10. Module 6 — Dispatch & Logistics Management", 1)

add_heading_styled(doc, "10.1 Dispatch Planning", 2)
dl_plan = [
    ("DL-001", "Group ready material by customer, destination, urgency for optimal truck loads", "High"),
    ("DL-002", "Inbound: multi-point pickup (one truck, multiple mills in same region)", "High"),
    ("DL-003", "Outbound: multi-drop delivery planning (one truck, multiple customers)", "High"),
    ("DL-004", "Flag truck-load readiness vs. need for combining with other orders", "High"),
    ("DL-005", "Urgency-based dispatch: urgent items go even if truck isn't full", "High"),
    ("DL-006", "Generate loading plan document for mill/godown", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], dl_plan, col_widths=[2, 12, 2])

add_heading_styled(doc, "10.2 Bill Interception Workflow", 2)
add_para(doc, (
    "Monit's unique workflow: mill sends goods with mill's invoice. Monit's field person intercepts "
    "the truck en route, swaps the mill invoice with Monit's own customer invoice before delivery."
), size=9)
dl_bill = [
    ("DL-010", "Record mill invoice details upon dispatch: invoice #, lorry #, driver, quantities", "High"),
    ("DL-011", "Auto-generate Monit's customer invoice (via Tally) for the interception", "High"),
    ("DL-012", "Assign field person for interception: which truck, where, printable invoices", "High"),
    ("DL-013", "Field person confirms interception: mill invoice collected, Monit invoice given", "Medium"),
    ("DL-014", "Mandatory interception for direct delivery (mill → customer) shipments", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], dl_bill, col_widths=[2, 12, 2])

add_heading_styled(doc, "10.3 Transportation & Compliance", 2)
dl_transport = [
    ("DL-020", "Transporter master: name, vehicles, rate/ton-km, service areas", "Medium"),
    ("DL-021", "Track freight costs per dispatch (inbound, outbound, direct delivery)", "Medium"),
    ("DL-022", "Freight allocation to specific SOs for profit calculation", "Medium"),
    ("DL-023", "E-Way Bill generation / portal integration for GST compliance", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], dl_transport, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 11 — BILLING & TALLY
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "11. Module 7 — Billing, Invoicing & Tally Integration", 1)

add_heading_styled(doc, "11.1 Invoice Generation", 2)
bl_inv = [
    ("BL-001", "Generate Sales Invoices from confirmed SOs and dispatched quantities", "High"),
    ("BL-002", "Correct firm selection: Agency (ITC) / Associates (NR/others) — configurable mapping", "High"),
    ("BL-003", "Advance billing support: invoice before physical delivery when confirmed", "High"),
    ("BL-004", "Include: GSTIN, HSN codes, GST rates, paper-specific details (type, GSM, size, qty)", "High"),
    ("BL-005", "Debit Notes / Credit Notes for rate adjustments, returns, quality claims, shortages", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], bl_inv, col_widths=[2, 12, 2])

add_heading_styled(doc, "11.2 Tally Synchronization", 2)
add_para(doc, "CRITICAL INTEGRATION — Tally remains the accounting backbone.", bold=True, color=BRAND_RED, size=9)
bl_tally = [
    ("BL-010", "Two-way sync: Push invoices/DN/CN to Tally; Pull payments/balances from Tally", "Critical"),
    ("BL-011", "Near-real-time sync (< 5 min) or configurable batch sync", "High"),
    ("BL-012", "Pull purchase bill data from Tally to update procurement records", "High"),
    ("BL-013", "Continuous payment status sync — drives approval workflow and AI chatbot", "Critical"),
    ("BL-014", "Handle firm bifurcation: Agency and Associates → separate Tally company files", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], bl_tally, col_widths=[2, 12, 2])

add_heading_styled(doc, "11.3 GST & Tax Compliance", 2)
bl_gst = [
    ("BL-020", "E-Invoice generation (mandatory above threshold turnover)", "High"),
    ("BL-021", "E-Way Bill generation for every dispatch above threshold", "High"),
    ("BL-022", "GSTR reconciliation — match purchase register with GSTR-2A/2B", "Medium"),
    ("BL-023", "TCS (Tax Collected at Source) on sales above ₹50 lakh/customer/year", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], bl_gst, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 12 — PAYMENT & CREDIT
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "12. Module 8 — Payment & Credit Management", 1)
add_para(doc, (
    "Paper trading operates on a unique financial model: mills require 100% advance payment, "
    "but dealers extend unsecured credit to customers. Managing this credit exposure is critical."
), bold=True, size=9)

pm_reqs = [
    ("PM-001", "Configurable credit limit per customer (set by owner)", "High"),
    ("PM-002", "Real-time credit utilization: limit - outstanding - pending invoices - pipeline orders", "Critical"),
    ("PM-003", "Credit health indicators on every interaction: limit, utilized, available, overdue", "High"),
    ("PM-004", "Temporary credit limit adjustments (increase/decrease) with date range and reason", "High"),
    ("PM-005", "Auto-categorize: Prompt, Slow (15-30d late), Risky (30+d late), Defaulter", "Medium"),
    ("PM-010", "Pull payment receipts from Tally — show all payments against invoices", "High"),
    ("PM-011", "PDC tracking: cheque #, date, amount, bank, status (received/deposited/cleared/bounced)", "Medium"),
    ("PM-012", "Automated payment reminders to sales team (not directly to customers)", "High"),
    ("PM-013", "Aging analysis per customer: current, 30d, 60d, 90d, 120+ days", "High"),
    ("PM-020", "Daily collection follow-up list per salesperson", "High"),
    ("PM-021", "Collection activity logging: called, visited, promised date, partial payment", "Medium"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], pm_reqs, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 13 — QUALITY & CLAIMS
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "13. Module 9 — Quality, Claims & Returns Management", 1)
qc_reqs = [
    ("QC-001", "Complaint registration: customer, invoice, paper details, complaint type, photos", "High"),
    ("QC-002", "Categorize: mill defect, sheeter defect, transit damage, customer handling", "High"),
    ("QC-003", "Generate claim document for mill with details, evidence, requested resolution", "High"),
    ("QC-004", "Claim lifecycle: raised → submitted → acknowledged → investigated → resolved", "Medium"),
    ("QC-005", "Customer returns: quantity, reason, condition, resolution (replacement/CN/restock)", "High"),
    ("QC-006", "Quality metrics per mill: complaint rate, claim acceptance rate, resolution time", "Medium"),
    ("QC-007", "Auto-generate weight shortage claims when discrepancy exceeds tolerance", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], qc_reqs, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 14 — REPORTING & ANALYTICS
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "14. Module 10 — Reporting, Analytics & Market Intelligence", 1)
add_para(doc, (
    "Owner's primary demand: \"I want auto-generated reports so I don't have to manually filter Excel.\" "
    "This module provides comprehensive automated reporting and market intelligence from system data."
), bold=True, size=9)

add_heading_styled(doc, "14.1 Sales Reports", 2)
rp_sales = [
    ("RP-001", "Sales Summary — by firm, mill, product, customer with period comparison", "Daily/Weekly/Monthly"),
    ("RP-002", "Salesperson Performance — orders, revenue, inquiries, conversion rate", "Weekly/Monthly"),
    ("RP-003", "Customer-Wise Sales — revenue, product mix, order size, frequency", "Monthly"),
    ("RP-004", "Mill-Wise Sales — product sales volume, margin analysis per mill", "Monthly"),
    ("RP-005", "Product-Wise Sales — best-selling types, GSMs, sizes; trend analysis", "Monthly"),
]
styled_table(doc, ["ID", "Report", "Frequency"], rp_sales, col_widths=[2, 12, 3])

add_heading_styled(doc, "14.2 Procurement & Inventory Reports", 2)
rp_proc = [
    ("RP-010", "Pending Orders — open POs with status, dates, overdue flags", "Daily"),
    ("RP-011", "Mill Performance Scorecard — TAT adherence, weight accuracy, quality rating", "Monthly"),
    ("RP-020", "Stock Position — by location, type, GSM, size with status breakdown", "Real-time/Daily"),
    ("RP-021", "Stock Aging — items >30, >60, >90 days; slow-moving identification", "Weekly"),
    ("RP-022", "Stock Turnover — days of inventory on hand per category", "Monthly"),
    ("RP-023", "Sheeter Output — reel consumption, sheets produced, waste %", "Daily/Weekly"),
]
styled_table(doc, ["ID", "Report", "Frequency"], rp_proc, col_widths=[2, 12, 3])

add_heading_styled(doc, "14.3 Financial & Follow-Up Reports", 2)
rp_fin = [
    ("RP-030", "Receivables Aging — outstanding by customer, age-bucketed", "Daily"),
    ("RP-031", "Collection Report — payments received vs dues", "Daily/Weekly"),
    ("RP-032", "Margin Analysis — per-transaction, per-customer, per-mill", "Monthly"),
    ("RP-040", "Mill-Wise Customer Follow-Up — who hasn't been contacted, what they buy, last order", "Daily"),
    ("RP-041", "Inquiry Response Time — by channel and salesperson", "Weekly"),
    ("RP-042", "Lost Inquiry Analysis — non-conversion reasons, patterns, competitor indicators", "Monthly"),
]
styled_table(doc, ["ID", "Report", "Frequency"], rp_fin, col_widths=[2, 12, 3])

add_heading_styled(doc, "14.4 Market Intelligence (from App/Chatbot Data)", 2)
rp_market = [
    ("RP-050", "App search patterns: most-searched paper types, GSMs, sizes → demand signals", "High"),
    ("RP-051", "Unmet demand: searches for products not in stock → stocking opportunity", "High"),
    ("RP-052", "Inquiry-to-order drop-off analysis: searched → available → didn't order. Why?", "High"),
    ("RP-053", "Competitor indicators: regular customer volume drop, search-but-don't-buy patterns", "Medium"),
    ("RP-054", "Demand forecasting: predict 2-4 week demand by product category", "Medium"),
]
styled_table(doc, ["ID", "Intelligence", "Priority"], rp_market, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 15 — CRM
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "15. Module 11 — CRM & Customer Relationship Management", 1)
crm_reqs = [
    ("CR-001", "Customer master: company, GST, contacts, addresses, communication preferences", "High"),
    ("CR-002", "Primary salesperson assignment — all notifications go to mapped salesperson first", "High"),
    ("CR-003", "Customer tiering: Platinum, Gold, Silver, Bronze with tier-based SLAs", "Medium"),
    ("CR-004", "360° interaction timeline: calls, visits, inquiries, orders, complaints", "High"),
    ("CR-005", "Auto follow-up: 'Customer X hasn't ordered in Y days (usually orders every Z days)'", "High"),
    ("CR-006", "Sales visit planning with nearby customer suggestions", "Low"),
    ("CR-007", "Sample tracking: what sent, when, follow-up date, outcome (ordered/rejected/pending)", "Medium"),
    ("CR-008", "Product preference map: per customer — types, GSMs, sizes, mills they regularly buy", "High"),
]
styled_table(doc, ["ID", "Requirement", "Priority"], crm_reqs, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 16 — CONFIGURATION & MASTER DATA
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "16. Module 12 — Configuration & Master Data", 1)

add_heading_styled(doc, "16.1 Master Data", 2)
md_list = [
    ("MD-001", "Paper Type Master", "Type, sub-type, HSN code, GST rate"),
    ("MD-002", "GSM Master", "Available GSMs per paper type"),
    ("MD-003", "Size Master", "Standard sheet sizes, reel deckle widths"),
    ("MD-004", "Mill Master", "Name, location, products, TAT, payment terms"),
    ("MD-005", "Customer Master", "See CRM module (CR-001)"),
    ("MD-006", "Transporter Master", "Name, vehicles, rates, service areas"),
    ("MD-007", "Location/Godown Master", "Address, capacity, type"),
    ("MD-008", "Box Type Master", "20-40 carton types for planning engine"),
    ("MD-009", "Die Master", "Existing dies — dimensions, customer, usage count"),
    ("MD-010", "Firm Master", "Agency & Associates: GST, bank accounts, Tally mapping"),
]
styled_table(doc, ["ID", "Master", "Key Fields"], md_list, col_widths=[2, 5, 10])

add_heading_styled(doc, "16.2 Configurable Settings", 2)
cf_list = [
    ("CF-001", "Credit limit defaults by customer tier"),
    ("CF-002", "Approval escalation thresholds (credit utilization %)"),
    ("CF-003", "Stock lock timeout duration (e.g., 2 hours)"),
    ("CF-004", "Inquiry response SLA per channel"),
    ("CF-005", "Weight tolerance % for reconciliation"),
    ("CF-006", "Mill-to-Firm billing mapping"),
    ("CF-007", "Minimum margin rules per product category"),
    ("CF-008", "TAT settings per mill/product/category"),
    ("CF-009", "Customer qualification flags for auto-processing"),
    ("CF-010", "Stock visibility rules (show/hide/restrict)"),
    ("CF-011", "Notification message templates (WhatsApp, email, app)"),
    ("CF-012", "Chatbot escalation trigger rules"),
]
styled_table(doc, ["ID", "Setting"], cf_list, col_widths=[2, 15])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 17 — NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "17. Non-Functional Requirements", 1)

add_heading_styled(doc, "17.1 Performance", 2)
nf_perf = [
    ("NF-001", "Stock availability check: < 2 seconds"),
    ("NF-002", "AI chatbot response: < 5 seconds"),
    ("NF-003", "Dashboard / reports loading: < 3 seconds"),
    ("NF-004", "Support: 50 concurrent internal users + 500 customer app users"),
    ("NF-005", "Tally sync: < 5 minutes of transaction"),
]
styled_table(doc, ["ID", "Requirement"], nf_perf, col_widths=[2, 15])

add_heading_styled(doc, "17.2 Availability, Security & Scalability", 2)
nf_other = [
    ("NF-010", "99.5% uptime during business hours (Mon–Sat 9:30AM–7PM IST)"),
    ("NF-011", "Daily data backup with 30-day retention"),
    ("NF-012", "Graceful degradation if Tally sync is down"),
    ("NF-020", "Authenticated access with OTP for sensitive actions"),
    ("NF-021", "Role-based access control (RBAC)"),
    ("NF-022", "HTTPS/TLS encryption for all API communication"),
    ("NF-023", "Customer financial data never visible to customers"),
    ("NF-024", "Audit logs for all sensitive operations"),
    ("NF-030", "Scale from ~100 to 500+ active customers over 3 years"),
    ("NF-031", "Mills, products, locations addable via config — no code changes"),
    ("NF-040", "Hindi + English support for internal UI; Hinglish for chatbot"),
    ("NF-041", "Simple customer app: max 3 taps to check availability"),
    ("NF-050", "Data migration: existing Excel, Google Sheets, Tally historical data"),
]
styled_table(doc, ["ID", "Requirement"], nf_other, col_widths=[2, 15])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 18 — BUSINESS RULES ENGINE
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "18. Business Rules Engine", 1)
add_para(doc, (
    "The owner explicitly wants business rules to be configurable, not hardcoded: "
    "\"The system should be the expert, not the person.\" Rules drive automated decisions across all modules."
))

add_heading_styled(doc, "18.1 Credit & Payment Rules", 2)
br_credit = [
    ("BR-001", "Credit utilization > 100%", "Block order, escalate to Owner"),
    ("BR-002", "Credit utilization > 90%", "Flag for owner review, two-step approval"),
    ("BR-003", "Overdue > 30 days", "Warn salesperson, require owner approval"),
    ("BR-004", "Overdue > 60 days", "Auto-flag as 'Watch'; chatbot says 'let me check'"),
    ("BR-005", "Customer 'Blocked'", "No orders; chatbot says 'connect with team'"),
]
styled_table(doc, ["Rule ID", "Condition", "Action"], br_credit, col_widths=[2, 6, 9])

add_heading_styled(doc, "18.2 Stock & Availability Rules", 2)
br_stock = [
    ("BR-010", "Stock available + rules pass", "Show 'Available' immediately"),
    ("BR-011", "Stock available + rules fail (soft)", "'Checking availability…' → internal escalation"),
    ("BR-012", "Stock available + rules fail (hard)", "'Connecting with team' → escalation"),
    ("BR-013", "Stock not available", "'Not Available' with alternatives"),
    ("BR-014", "Partial stock", "'Partially Available' with timeline"),
]
styled_table(doc, ["Rule ID", "Condition", "Action"], br_stock, col_widths=[2, 6, 9])

add_heading_styled(doc, "18.3 Situational Overrides", 2)
br_override = [
    ("BR-040", "'Mill Pressure Mode' — temporarily relax credit limits by configurable % for set duration", "High"),
    ("BR-041", "All overrides logged with reason; auto-revert after expiry", "High"),
    ("BR-042", "Ad-hoc customer-specific rules: e.g., 'Allow 150% credit for 7 days for customer X'", "Medium"),
]
styled_table(doc, ["Rule ID", "Description", "Priority"], br_override, col_widths=[2, 12, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 19 — INTEGRATION REQUIREMENTS
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "19. Integration Requirements", 1)

integ_table = [
    ("1", "Tally ERP", "Bi-directional", "Invoices, payments, ledger, purchase bills, DN/CN", "Critical"),
    ("2", "Planning Engine", "Outbound", "Box dimensions → optimal sheet sizes/plans", "High"),
    ("3", "WhatsApp Biz API", "Bi-directional", "Customer messages, chatbot, notifications", "High"),
    ("4", "E-Way Bill Portal", "Outbound", "Invoice details → E-Way Bill generation", "High"),
    ("5", "E-Invoice Portal", "Outbound", "Invoice details → IRN generation", "High"),
    ("6", "Email (SMTP)", "Outbound", "POs to mills, confirmations, alerts", "Medium"),
    ("7", "SMS Gateway", "Outbound", "OTP, critical alerts", "Low"),
    ("8", "ITC Portal (Future)", "Inbound", "Order status, dispatch details", "Future"),
]
styled_table(doc, ["#", "System", "Direction", "Data Exchanged", "Priority"], integ_table, col_widths=[1, 3, 3, 7, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 20 — GROWTH AFTER AUTOMATION (THE KEY NEW SECTION)
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "20. Growth After Automation — Impact Analysis", 1)

add_para(doc, (
    "This section presents the projected operational and business impact of implementing the "
    "Integrated Paper Trading Management System. The analysis compares current manual operations "
    "with the automated system across six key dimensions."
), size=10)

# ── 20.1 Executive Summary ────────────────────────────────────
add_heading_styled(doc, "20.1 Executive Impact Summary", 2)
impact_summary = [
    ("Manpower Optimization", "16 people across functions", "5–6 people (same or more output)", "~60% reduction"),
    ("Inquiry Response Time", "30–120 minutes (manual)", "1–5 minutes (automated)", "~95% faster"),
    ("Double Stock Booking", "10–15 incidents/month", "0 incidents (system-locked)", "100% eliminated"),
    ("Daily MIS / Reports", "2+ hours manual Excel work", "Instant auto-generated", "100% automated"),
    ("Inquiry-to-Order Conversion", "~20% conversion rate", "~55% (faster response, AI suggestions)", "2.75× improvement"),
    ("Customer Reach (24/7)", "Business hours only", "24/7 via app & chatbot", "3× availability window"),
    ("Revenue Growth (2-Year)", "Baseline", "+75% projected (compounding)", "Significant growth"),
    ("Payment Overdue Incidents", "Frequent (manual checks)", "Rare (system auto-blocks)", "~90% reduction"),
]
styled_table(doc,
    ["Metric", "Before Automation", "After Automation", "Impact"],
    impact_summary,
    col_widths=[4, 4.5, 5, 3.5],
    header_color=BRAND_GREEN
)

doc.add_paragraph()

# ── 20.2 Manpower Optimization Chart ──────────────────────────
add_heading_styled(doc, "20.2 Manpower Optimization", 2)
add_para(doc, (
    "The most immediate impact is the reduction in people-dependency. The system automates "
    "repetitive tasks (stock updates, report generation, follow-up tracking), allowing the business "
    "to handle 2–3× volume with the same or fewer people. New hires don't need to be 'experts' — "
    "the system guides them."
))
doc.add_picture(CHART_MANPOWER, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

manpower_detail = [
    ("Sales Follow-up", "4 people doing manual Excel filtering, phone calls, WhatsApp coordination", "2 people: system auto-generates follow-up lists, AI chatbot handles routine inquiries"),
    ("Procurement Tracking", "3 people tracking orders via Excel, Word POs, email follow-ups", "1 person: system tracks all POs, auto-alerts on delays, generates PO PDFs"),
    ("Billing & Invoicing", "3 people: manual Word/Tally entry, bill interception coordination", "1 person: auto-generated invoices synced with Tally, digital interception workflow"),
    ("Stock Management", "2 people: daily Google Sheet updates, manual stock checks", "1 person: real-time system updates, automated transit tracking"),
    ("Dispatch Planning", "2 people: manual truck load planning, phone coordination", "1 person: system suggests optimal loads, digital loading plans"),
    ("Reporting & MIS", "2 people: daily Excel report compilation, manual filtering", "0 people: fully automated, scheduled, real-time dashboards"),
]
styled_table(doc, ["Function", "Before (Manual)", "After (Automated)"], manpower_detail, col_widths=[3, 7, 7])

doc.add_paragraph()

# ── 20.3 Response Time Chart ──────────────────────────────────
add_heading_styled(doc, "20.3 Response Time Improvement", 2)
add_para(doc, (
    "Speed of response is the #1 factor in converting inquiries to orders. Currently, customers "
    "wait 30–120 minutes for availability confirmation. With automation, this drops to under 5 minutes — "
    "often instant via the AI chatbot. This directly prevents customers from sourcing elsewhere."
))
doc.add_picture(CHART_RESPONSE, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── 20.4 Error Elimination Chart ──────────────────────────────
add_heading_styled(doc, "20.4 Operational Error Elimination", 2)
add_para(doc, (
    "The most painful current issue: one salesman sells stock already allocated to another customer's order. "
    "This damages customer relationships permanently. The real-time stock reservation system eliminates this "
    "entirely. Other errors — wrong firm billing, missed follow-ups, credit limit violations — are "
    "similarly addressed by system-enforced rules."
))
doc.add_picture(CHART_ERRORS, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── 20.5 Inquiry Conversion Chart ─────────────────────────────
add_heading_styled(doc, "20.5 Inquiry-to-Delivery Conversion Improvement", 2)
add_para(doc, (
    "Currently, out of every 100 inquiries, only ~20 convert to confirmed orders — primarily due to "
    "slow response times, inability to suggest alternatives, and lack of systematic follow-up. "
    "The AI chatbot with instant availability, proactive suggestions, and the planning engine can "
    "more than double this conversion rate."
))
doc.add_picture(CHART_CONVERSION, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── 20.6 Daily Operations Time Chart ──────────────────────────
add_heading_styled(doc, "20.6 Daily Operational Time Savings", 2)
add_para(doc, (
    "Current daily operations consume ~14 person-hours on repetitive tasks that can be automated. "
    "After implementation, this drops to ~2.5 person-hours — freeing the team to focus on "
    "relationship building, new customer acquisition, and strategic decisions."
))
doc.add_picture(CHART_DAILY_OPS, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── 20.7 Revenue Growth Projection ────────────────────────────
add_heading_styled(doc, "20.7 Projected Revenue Growth", 2)
add_para(doc, (
    "The compound effect of faster response, higher conversion, 24/7 availability, intelligent "
    "cross-selling (AI suggesting alternatives, stock lots, reel-cuts), and reduced operational friction "
    "is projected to drive significant revenue growth. Conservative estimate: 75% growth over 2 years "
    "from the same customer base, with additional upside from new customers attracted by the app/chatbot."
))
doc.add_picture(CHART_REVENUE, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── 20.8 Key Growth Drivers ───────────────────────────────────
add_heading_styled(doc, "20.8 Key Growth Drivers After Automation", 2)

growth_drivers = [
    ("Zero Double Bookings", "Real-time stock locking eliminates the #1 customer sentiment destroyer. Every salesman sees live, accurate stock — including what's booked, in transit, and available.", "Customer retention, trust"),
    ("24/7 Customer Access", "Customers can check availability, get AI suggestions, and place inquiries anytime — not just during 9:30AM–7PM. Night-time and early-morning inquiries are captured automatically.", "+30% inquiry volume"),
    ("AI Cross-Selling", "Chatbot proactively suggests alternatives: different GSM, different mill, reel-cut option, stock lot deals. Converts 'not available' into 'here's what we can do.'", "+15% order value"),
    ("Planning Engine Advantage", "No competitor offers carton-dimension-to-paper-size optimization. This is a unique differentiator that locks in customers (especially Vijay Shri-type printers).", "Competitive moat"),
    ("Market Intelligence", "Every app search, every inquiry, every 'search but didn't order' — feeds into demand analytics. Know what the market wants before competitors do.", "Strategic advantage"),
    ("Reduced Dependency on Experts", "New staff don't need years of product knowledge — the system guides decisions. Scale the team without quality loss.", "Scalability"),
    ("Mill Relationship Optimization", "System tracks TAT adherence, quality, and fill rates per mill. Data-driven decisions on which mills to prioritize.", "Better procurement terms"),
    ("Payment Discipline", "Automated credit checks prevent overexposure. Collection follow-ups are systematic, not forgotten.", "Cash flow improvement"),
]
styled_table(doc,
    ["Growth Driver", "How It Works", "Impact"],
    growth_drivers,
    col_widths=[3.5, 10, 3.5],
    header_color=BRAND_ACCENT
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 21 — PHASED ROLLOUT
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "21. Phased Rollout Plan", 1)

add_heading_styled(doc, "Phase 1 — Foundation & Procurement (Months 1–3)", 2)
add_para(doc, "Focus: Replace Excel/Word-based procurement with a structured system.", bold=True)
p1 = [
    ("Master Data Setup", "Mill, product, customer, location masters"),
    ("Purchase Order Module", "Order form, PO generation (PDF), email to mill"),
    ("Pending Order Tracking", "Status tracking, delivery dates, follow-up alerts"),
    ("Material Receipt", "Transit tracking, godown receipt, weight reconciliation"),
    ("Basic Inventory", "Multi-location stock with 5 statuses"),
    ("Tally Bridge (Phase 1)", "Pull payment/outstanding data from Tally"),
    ("Dashboards", "Procurement dashboard, pending orders, stock position"),
]
styled_table(doc, ["Deliverable", "Description"], p1, col_widths=[5, 12])

add_heading_styled(doc, "Phase 2 — Sales & Internal Ops (Months 3–5)", 2)
add_para(doc, "Focus: Digitize sales workflow, stock reservation, approval engine.", bold=True)
p2 = [
    ("Inquiry Management", "Capture, assignment, tracking"),
    ("Sales Order Workflow", "SO creation, confirmation, fulfillment tracking"),
    ("Stock Reservation", "Real-time soft-locking (salesman conflict fix)"),
    ("Approval Workflow", "Rule-based with multi-level escalation"),
    ("Dispatch Planning", "Load planning, bill interception, delivery tracking"),
    ("Tally Bridge (Phase 2)", "Push invoices, pull updated payments"),
    ("Sales Reports", "Performance, follow-up, conversion reports"),
]
styled_table(doc, ["Deliverable", "Description"], p2, col_widths=[5, 12])

add_heading_styled(doc, "Phase 3 — Customer App & AI Chatbot (Months 5–8)", 2)
add_para(doc, "Focus: Launch the 'game changer' customer-facing experience.", bold=True)
p3 = [
    ("Mobile App", "Registration, search, availability, order, tracking"),
    ("AI Chatbot (App)", "Conversational availability, alternatives, inquiry/order flow"),
    ("WhatsApp Chatbot", "Same capabilities via WhatsApp Business API"),
    ("Planning Engine", "Carton dimensions → optimal sheet size recommendation"),
    ("Hidden Rule Engine", "Silent credit checks with diplomatic responses"),
    ("Market Intelligence", "Search analytics, demand tracking"),
]
styled_table(doc, ["Deliverable", "Description"], p3, col_widths=[5, 12])

add_heading_styled(doc, "Phase 4 — Advanced Features (Months 8–12)", 2)
add_para(doc, "Focus: Refine, optimize, and add advanced capabilities.", bold=True)
p4 = [
    ("Credit Management", "Full lifecycle, PDC, collection follow-up"),
    ("Quality & Claims", "Complaint registration, mill claims, returns"),
    ("Advanced Reporting", "Forecasting, mill scorecards, margin analysis"),
    ("CRM Enhancements", "Visit planning, sample tracking, journey analytics"),
    ("AI Improvements", "Learning from data, better recommendations"),
    ("GST Compliance", "E-Invoice, E-Way Bill automation, GSTR reconciliation"),
]
styled_table(doc, ["Deliverable", "Description"], p4, col_widths=[5, 12])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 22 — APPENDICES
# ══════════════════════════════════════════════════════════════
add_heading_styled(doc, "22. Appendices", 1)

add_heading_styled(doc, "Appendix A — Mill Portfolio", 2)
mill_portfolio = [
    ("ITC Limited (PSPD)", "FBB, SBS, Coated, Packaging Board", "Monit Paper Sales Agency", "3–7 days"),
    ("NR Agarwal Industries", "Duplex Board (GB/WB), Kraft", "Monit Paper Associates", "4–10 days"),
    ("Gayatri Shakti Paper", "Duplex Board, Kraft", "Associates", "5–12 days"),
    ("Emami Paper Mills", "Writing/Printing, Packaging", "Associates", "7–15 days"),
    ("Balkrishna Paper Mills", "Duplex Board", "Associates", "5–10 days"),
    ("JK Paper Ltd", "FBB, Writing/Printing", "Agency / Associates", "5–15 days"),
    ("Others (6-7 mills)", "Various", "As configured", "Variable"),
]
styled_table(doc, ["Mill", "Products", "Billing Firm", "Typical TAT"], mill_portfolio, col_widths=[4, 5, 4.5, 3])

add_heading_styled(doc, "Appendix B — Product Catalog Structure", 2)
prod_cat = [
    ("Duplex Board", "Grey Back (GBD), White Back (WBD)", "200–450 GSM", "23x36, 25x36, 28x36"),
    ("FBB", "Standard, Premium", "210–400 GSM", "23x36, 25x36, 22x28"),
    ("SBS", "Standard", "200–350 GSM", "Various"),
    ("Kraft Paper", "Brown, White", "80–300 GSM", "Reel and sheet"),
    ("Coated Paper", "C1S, C2S", "80–300 GSM", "Various"),
]
styled_table(doc, ["Paper Type", "Sub-Types", "GSM Range", "Common Sizes"], prod_cat, col_widths=[4, 4, 4, 5])

add_heading_styled(doc, "Appendix C — Additional Paper Trading Activities", 2)
add_para(doc, (
    "The following 25 activities are standard in the paper trading industry. "
    "All have been incorporated into the relevant modules of this SRS:"
), size=9)
activities = [
    ("1", "Weight Reconciliation (mill vs actual)", "Procurement (PR-033)"),
    ("2", "Quality Claims & Returns", "Quality Module (Section 13)"),
    ("3", "Sample Management", "CRM (CR-007)"),
    ("4", "Debit Note / Credit Note Management", "Billing (BL-005)"),
    ("5", "Mill Commission & Incentive Tracking", "Procurement (PR-006)"),
    ("6", "Freight & Transport Cost Management", "Dispatch (DL-020–022)"),
    ("7", "GST / E-Way Bill / E-Invoice Compliance", "Billing (BL-020–023)"),
    ("8", "Payment Collection & PDC Management", "Payment (Module 12)"),
    ("9", "Customer Segmentation & Tiering", "CRM (CR-003)"),
    ("10", "Price List & Margin Management", "Sales (SL-010–015)"),
    ("11", "Sheeter Trim/Waste Management", "Inventory (IN-020–024)"),
    ("12", "Multi-Location Godown Management", "Inventory (IN-001–004)"),
    ("13", "Transit & Stock Insurance Claims", "Quality (QC-002)"),
    ("14", "Seasonal Demand Forecasting", "Reporting (RP-054)"),
    ("15", "Competitor Price Intelligence", "Reporting (RP-053)"),
    ("16", "Sales Team Visit/Activity Tracking", "CRM (CR-006)"),
    ("17", "Mill Scheme Pass-through", "Procurement (PR-006)"),
    ("18", "Stock Lot / Seconds Management", "Inventory (IN-030–032)"),
    ("19", "Reel Inventory & Conversion Tracking", "Inventory (IN-020–024)"),
    ("20", "Customer Machine Profile Management", "Customer App (CA-050–052)"),
    ("21", "Die Library Management", "Config (MD-009)"),
    ("22", "Inter-Godown Stock Transfers", "Inventory (IN-003)"),
    ("23", "Mill Run Broadcast & Proactive Sales", "Sales (SL-006)"),
    ("24", "Advance Payment to Mills (100% advance model)", "Payment (PM-014)"),
    ("25", "Unsecured Credit Risk Management", "Payment (PM-001–005)"),
]
styled_table(doc, ["#", "Activity", "Covered In"], activities, col_widths=[1, 8, 8])

doc.add_paragraph()
doc.add_paragraph()

# ── FOOTER ─────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Document —")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string(BRAND_DARK.lstrip("#"))
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "This document is subject to review and approval by the Monit Paper Agency management team.\n"
    "Requirements may be refined during the detailed design phase."
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor.from_string("666666")
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL — Monit Paper Sales Agency / Monit Paper Associates")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor.from_string(BRAND_RED.lstrip("#"))
run.bold = True

# ── SAVE ───────────────────────────────────────────────────────
output_path = os.path.join(OUTPUT_DIR, "SRS_Monit_Paper_Agency.docx")
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")
print("Done!")
